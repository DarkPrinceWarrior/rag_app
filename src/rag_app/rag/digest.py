"""Выжимка/история чата (§ 5 / § 6): executive-сводка диалога + транскрипт +
дедуплицированные источники → Markdown / DOCX для экспорта."""

from __future__ import annotations

import io
from typing import Any

from docx import Document as Docx
from openai import AsyncOpenAI

from rag_app.config import settings


def _src_line(i: int, c: dict[str, Any]) -> str:
    page = f", стр. {c['page_start'] + 1}" if c.get("page_start") is not None else ""
    head = f" — {c['heading_path']}" if c.get("heading_path") else ""
    return f"{i}. {c.get('filename', '')}{page}{head}"


async def session_digest(client: AsyncOpenAI, title: str, messages: list[Any]) -> dict[str, Any]:
    """LLM-выжимка + транскрипт + источники (дедуп по документу/странице/разделу)."""
    convo = "\n".join(f"{m.role}: {m.content[:1200]}" for m in messages)
    summary = ""
    if convo.strip():
        try:
            resp = await client.chat.completions.create(
                model=settings.llm_model,
                messages=[
                    {
                        "role": "user",
                        "content": (
                            f"Диалог пользователя с технической документацией:\n{convo[:12000]}\n\n"
                            "Сделай краткую выжимку на русском: 4–8 пунктов с ключевыми выводами, "
                            "числами и условиями из ответов. Только пункты, без вступлений."
                        ),
                    }
                ],
                temperature=0.2,
                max_tokens=600,
                extra_body={"chat_template_kwargs": {"enable_thinking": False}},
            )
            summary = (resp.choices[0].message.content or "").strip()
        except Exception:
            summary = ""

    seen: dict[tuple, dict] = {}
    for m in messages:
        for c in m.citations or []:
            key = (c.get("document_id"), c.get("page_start"), c.get("heading_path"))
            seen.setdefault(key, c)

    return {
        "title": title,
        "summary": summary,
        "messages": [{"role": m.role, "content": m.content} for m in messages],
        "sources": list(seen.values()),
    }


def render_md(d: dict[str, Any]) -> str:
    lines = [f"# {d['title']}", ""]
    if d["summary"]:
        lines += ["## Выжимка", "", d["summary"], ""]
    lines += ["## Диалог", ""]
    for m in d["messages"]:
        who = "Вопрос" if m["role"] == "user" else "Ответ"
        lines += [f"**{who}:** {m['content']}", ""]
    if d["sources"]:
        lines += ["## Источники", ""]
        lines += [_src_line(i, c) for i, c in enumerate(d["sources"], 1)]
    return "\n".join(lines)


def render_docx(d: dict[str, Any]) -> bytes:
    doc = Docx()
    doc.add_heading(d["title"], level=1)
    if d["summary"]:
        doc.add_heading("Выжимка", level=2)
        for line in d["summary"].splitlines():
            if line.strip():
                doc.add_paragraph(line.strip())
    doc.add_heading("Диалог", level=2)
    for m in d["messages"]:
        who = "Вопрос" if m["role"] == "user" else "Ответ"
        p = doc.add_paragraph()
        p.add_run(f"{who}: ").bold = True
        p.add_run(m["content"])
    if d["sources"]:
        doc.add_heading("Источники", level=2)
        for i, c in enumerate(d["sources"], 1):
            doc.add_paragraph(_src_line(i, c))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
