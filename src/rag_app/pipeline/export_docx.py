"""Сборка редактируемого DOCX из переведённых сегментов (python-docx)."""

from __future__ import annotations

import io
import re

from docx import Document as DocxDocument
from docx.shared import Inches, Pt

from rag_app.db.models import Segment, SegmentKind

# Символы, недопустимые в XML 1.0 (python-docx/lxml роняет «All strings must be
# XML compatible … no NULL bytes»). Источник — OCR'нутые txt/сканы с битыми
# control-байтами. Чистим перед записью в DOCX.
_XML_INVALID = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\ud800-\udfff￾￿]")


def xml_safe(s: str | None) -> str:
    return _XML_INVALID.sub("", s) if s else (s or "")


def _seg_text(seg: Segment) -> str:
    return xml_safe(seg.translated_text if seg.translated_text is not None else seg.source_text)


# Инлайн-разметка из сегментов (MinerU/paddle дают **жирный** / *италик*). В DOCX
# превращаем в настоящие run'ы, иначе LibreOffice показывает буквальные «звёздочки».
_MD_SPLIT = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def _strip_md(s: str) -> str:
    """Убрать парные ** / * (оставить содержимое) — для заголовков и ячеек таблиц."""
    return re.sub(r"\*{1,2}([^*]+)\*{1,2}", r"\1", s)


def _add_md_runs(paragraph, text: str) -> None:
    """Добавить текст в абзац, разворачивая **жирный** / *италик* в run'ы."""
    for part in _MD_SPLIT.split(text):
        if not part:
            continue
        if len(part) > 4 and part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif len(part) > 2 and part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


# Инлайн-формулы MinerU/paddle приходят сырым LaTeX ($m \times n$, $x_{3}$,
# $\mathbf{s}$). DOCX/LibreOffice его не рендерит → показывал бы доллары и
# бэкслеши «кашей». Приводим к читаемому юникоду (порт web/lib/cleanMath.ts:
# $…$ снять, \times→×, ^{3}→³, команды/скобки убрать). Не настоящая математика,
# но читаемо. Дисплейные формулы (kind=equation) сложные — приблизит, не более.
_SUP = {"0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴", "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹"}


def _latex_to_text(s: str) -> str:
    if not s or not any(c in s for c in "$\\^{"):
        return s
    s = re.sub(r"\$([^$]*)\$", r"\1", s)  # снять $…$
    s = s.replace("\\circ", "°").replace("\\complement", "C")
    s = s.replace("\\times", "×").replace("\\cdot", "·").replace("\\pm", "±")
    s = re.sub(r"\\[,;: ]", " ", s)  # \, \; \: тонкие пробелы
    s = re.sub(r"\^\s*\{?\s*([0-9])\s*\}?", lambda m: _SUP.get(m.group(1), m.group(1)), s)  # ^{3}→³
    s = re.sub(r"\^\s*\{([^}]*)\}", r"\1", s)  # прочие ^{…}
    s = re.sub(r"\\[a-zA-Z]+", "", s)  # прочие \команды (\mathbf, \boldsymbol, …)
    s = s.replace("{", "").replace("}", "")
    return re.sub(r"[ \t]{2,}", " ", s).strip()


def build_docx(
    filename: str, segments: list[Segment], images: dict[str, bytes] | None = None
) -> bytes:
    """DOCX из переведённых сегментов. images: {img_s3 → байты} для встраивания
    рисунков (worker качает их из bucket_artifacts; без них — текстовая заглушка)."""
    doc = DocxDocument()
    doc.core_properties.title = filename

    for seg in segments:
        if seg.kind == SegmentKind.heading:
            level = min(max(seg.heading_level or 1, 1), 6)
            doc.add_heading(_strip_md(_latex_to_text(_seg_text(seg))), level=level)

        elif seg.kind == SegmentKind.paragraph:
            _add_md_runs(doc.add_paragraph(), _latex_to_text(_seg_text(seg)))

        elif seg.kind == SegmentKind.equation:
            p = doc.add_paragraph()
            run = p.add_run(_latex_to_text(xml_safe(seg.source_text)))
            run.italic = True
            run.font.size = Pt(10)

        elif seg.kind == SegmentKind.image:
            caption = _strip_md(_latex_to_text(_seg_text(seg))).strip()
            key = (seg.meta or {}).get("img_s3")
            blob = images.get(key) if images and key else None
            placed = False
            if blob:
                try:
                    pic = doc.add_picture(io.BytesIO(blob))
                    maxw = Inches(6)
                    if pic.width > maxw:
                        pic.height = int(pic.height * maxw / pic.width)
                        pic.width = maxw
                    if caption:
                        doc.add_paragraph().add_run(caption).italic = True
                    placed = True
                except Exception:  # noqa: BLE001 — битый/неподдержанный формат → заглушка
                    placed = False
            if not placed:
                doc.add_paragraph().add_run(
                    f"[Рисунок]{' ' + caption if caption else ''}"
                ).italic = True

        elif seg.kind == SegmentKind.table:
            rows = seg.meta.get("table_rows_ru") or seg.meta.get("table_rows") or []
            caption = _latex_to_text(
                xml_safe((seg.meta.get("caption_ru") or seg.meta.get("caption") or "").strip())
            )
            if caption:
                p = doc.add_paragraph()
                p.add_run(caption).bold = True
            if rows:
                n_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=n_cols)
                table.style = "Table Grid"
                for i, row in enumerate(rows):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = _strip_md(_latex_to_text(xml_safe(cell)))
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
